"""Формирование отчёта Word (.docx) из данных прогресса."""

from __future__ import annotations

from datetime import datetime
from io import BytesIO

from docx import Document
from docx.shared import Cm, Pt

def _fmt(value: float) -> str:
    return f"{round(value, 2):.2f}".replace(".", ",")


def _meters(pk: float) -> str:
    total = round(pk * 100)
    return f"{total // 1000} км {total % 1000:03d} м"


def _covered(layer: dict) -> float:
    return sum(max(0.0, float(s["e"]) - float(s["s"])) for s in layer.get("segments", []))


def _stats(layer: dict, start: float, end: float) -> tuple[float, float, float]:
    covered = _covered(layer)
    route_length = max(0.0, end - start)
    excluded = min(route_length, max(0.0, float(layer.get("excludedPk", 0) or 0)))
    total = route_length
    raw_covered = covered
    covered = max(0.0, raw_covered - excluded)
    remaining = max(0.0, total - raw_covered)
    return covered, remaining, covered / total * 100 if total else 0.0


def _active_project(data: dict) -> tuple[str, dict]:
    if "projects" not in data:
        return "Основной объект", data
    project = next((p for p in data.get("projects", []) if p.get("id") == data.get("activeProjectId")), data["projects"][0])
    return project.get("name", "Объект"), project.get("state", {})


def _add_table(document: Document, headers: list[str], rows: list[list[str]]):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
    for values in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, values):
            cell.text = str(text)
    return table


def build_word_document(data: dict) -> bytes:
    project_name, state = _active_project(data)
    start = float(state.get("start", 0))
    end = float(state["total"])
    length = max(0.0, end - start)
    document = Document()
    section = document.sections[0]
    section.top_margin = section.bottom_margin = Cm(1.8)
    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    document.add_heading("Отчёт по прогрессу укладки слоёв", 0)
    document.add_paragraph(f"Объект: {project_name}")
    document.add_paragraph(f"Дата формирования: {datetime.now():%d.%m.%Y %H:%M}")
    document.add_heading("1. Общие сведения", level=1)
    _add_table(document, ["Показатель", "Значение"], [["Протяжённость трассы", f"ПК {_fmt(start)} — {_fmt(end)} ({_meters(length)})"], ["Количество слоёв", str(len(state.get("layers", [])))]] )

    document.add_heading("2. План и факт", level=1)
    summary_rows = []
    for layer in state.get("layers", []):
        covered, remaining, actual = _stats(layer, start, end)
        plan = float(layer.get("planPercent", 0))
        plan_date = str(layer.get("planDate", "") or "")
        if len(plan_date) == 10:
            plan_date = ".".join(reversed(plan_date.split("-")))
        else:
            plan_date = "—"
        summary_rows.append([layer.get("name", ""), f"{plan:.1f}%", plan_date, f"{actual:.1f}%", f"{actual-plan:+.1f}", f"{_fmt(covered)} ПК", f"{_fmt(remaining)} ПК"])
    _add_table(document, ["Слой", "План", "Дата", "Факт", "Отклонение, п.п.", "Готово", "Осталось"], summary_rows)

    document.add_heading("3. Участки, контроль и исполнители", level=1)
    for index, layer in enumerate(state.get("layers", []), 1):
        document.add_heading(f"3.{index}. {layer.get('name', '')}", level=2)
        covered, remaining, actual = _stats(layer, start, end)
        document.add_paragraph(f"Выполнено: {actual:.1f}% ({_meters(covered)}). Осталось: {_meters(remaining)}.")
        segments = sorted(layer.get("segments", []), key=lambda item: item["s"])
        rows = []
        for n, seg in enumerate(segments, 1):
            rows.append([str(n), _fmt(seg["s"]), _fmt(seg["e"]), _meters(seg["e"]-seg["s"]), seg.get("status", "Выполнено"), seg.get("responsible", ""), seg.get("date", ""), seg.get("quality", ""), seg.get("note", "")])
        if rows:
            _add_table(document, ["№", "От, ПК", "До, ПК", "Длина", "Статус", "Ответственный", "Дата", "Качество", "Комментарий"], rows)
        else:
            document.add_paragraph("Участки не указаны.")

    document.add_paragraph("Отчёт сформирован автоматически. 1 ПК = 100 м.")
    output = BytesIO()
    document.save(output)
    return output.getvalue()
