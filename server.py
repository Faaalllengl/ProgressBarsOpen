#!/usr/bin/env python3
"""Локальный сервер: открывает программу и сохраняет данные в road-progress-data.json."""

import json
import sys
import webbrowser
from datetime import datetime
from http.server import HTTPServer, SimpleHTTPRequestHandler
from io import BytesIO
from pathlib import Path

from export_word import build_word_document
from openpyxl import Workbook, load_workbook
from docx import Document
from docx.shared import Cm, Pt

DIR = Path(__file__).resolve().parent
DATA_FILE = DIR / "road-progress-data.json"
LOTKI_DATA_FILE = DIR / "lotki-data.json"
BACKUP_DIR = DIR / "backups"
PORT = 8765
MAX_BACKUPS = 50


def _write_json_file(path: Path, data: dict):
    text = json.dumps(data, ensure_ascii=False, indent=2) + "\n"
    tmp = path.with_suffix(path.suffix + ".tmp")
    tmp.write_text(text, encoding="utf-8")
    tmp.replace(path)


def _project_state(data):
    """Return an active project state from both legacy and v2 data formats."""
    if "projects" not in data:
        return data
    projects = data.get("projects", [])
    active_id = data.get("activeProjectId")
    project = next((p for p in projects if p.get("id") == active_id), None)
    return (project or projects[0]).get("state", {}) if projects else {}


def _validate_segment(seg, start, total, label):
    if not isinstance(seg, dict):
        raise ValueError(f"{label}: участок должен быть объектом")
    try:
        segment_start, end = float(seg["s"]), float(seg["e"])
    except (KeyError, TypeError, ValueError):
        raise ValueError(f"{label}: нужны числовые значения «от» и «до»")
    if not (start <= segment_start < end <= total):
        raise ValueError(f"{label}: участок должен быть в пределах {start} — {total} ПК, при этом «до» больше «от»")


def _validate_state(state):
    if not isinstance(state, dict) or not isinstance(state.get("layers"), list):
        raise ValueError("ожидается объект с полем layers")
    try:
        start = float(state.get("start", 0))
        total = float(state["total"])
    except (KeyError, TypeError, ValueError):
        raise ValueError("нужна положительная общая протяжённость")
    if start < 0 or total <= start:
        raise ValueError("общая протяжённость должна быть положительной")
    for layer in state["layers"]:
        if not isinstance(layer, dict) or not str(layer.get("name", "")).strip():
            raise ValueError("у каждого слоя должно быть название")
        try:
            excluded = float(layer.get("excludedPk", 0) or 0)
        except (TypeError, ValueError):
            raise ValueError(f"слой «{layer['name']}»: исключение из расчёта должно быть числом")
        if excluded < 0 or excluded > total - start:
            raise ValueError(f"слой «{layer['name']}»: исключение из расчёта должно быть от 0 до {total - start} ПК")
        for index, segment in enumerate(layer.get("segments", []), 1):
            _validate_segment(segment, start, total, f"слой «{layer['name']}», участок {index}")


def _validate_data(data):
    if not isinstance(data, dict):
        raise ValueError("ожидается объект данных")
    if "projects" in data:
        projects = data.get("projects")
        if not isinstance(projects, list) or not projects:
            raise ValueError("нужен хотя бы один объект")
        for project in projects:
            if not str(project.get("name", "")).strip():
                raise ValueError("у каждого объекта должно быть название")
            _validate_state(project.get("state", {}))
    else:
        _validate_state(data)


def _validate_lotki(data):
    if not isinstance(data, dict) or not isinstance(data.get("segments"), list):
        raise ValueError("ожидается объект с полем segments")
    for index, segment in enumerate(data["segments"], 1):
        if not isinstance(segment, dict):
            raise ValueError(f"лоток {index}: должен быть объектом")
        try:
            start = float(segment["start"])
            end = float(segment["end"])
        except (KeyError, TypeError, ValueError):
            raise ValueError(f"лоток {index}: нужны числовые значения «от» и «до»")
        if start < 0 or end <= start:
            raise ValueError(f"лоток {index}: значение «до» должно быть больше «от»")
        sides = segment.get("sides")
        if not isinstance(sides, list) or not sides or not all(side in ("left", "right") for side in sides):
            raise ValueError(f"лоток {index}: выберите левую и/или правую сторону")
        if not isinstance(segment.get("completed", False), bool):
            raise ValueError(f"лоток {index}: некорректный статус")


def _lotki_length(segment):
    base_length = max(0.0, float(segment["end"]) - float(segment["start"])) * 100
    return base_length * max(1, len(segment.get("sides", [])))


def _lotki_export_excel(data):
    workbook = Workbook()
    summary = workbook.active
    summary.title = "Сводка"
    segments = data.get("segments", [])
    total = sum(_lotki_length(item) for item in segments)
    done = sum(_lotki_length(item) for item in segments if item.get("completed"))
    percent = done / total * 100 if total else 0
    summary.append(["Показатель", "Значение"])
    summary.append(["Модуль", "Лотки"])
    summary.append(["Количество отрезков", len(segments)])
    summary.append(["Общая длина, ПК", total])
    summary.append(["Выполнено, ПК", done])
    summary.append(["Выполнено, %", round(percent, 1)])
    sheet = workbook.create_sheet("Лотки")
    sheet.append(["№", "От ПК", "До ПК", "Сторона", "Длина, ПК", "Длина, м", "Статус"])
    for index, item in enumerate(segments, 1):
        sides = []
        if "left" in item.get("sides", []): sides.append("Слева")
        if "right" in item.get("sides", []): sides.append("Справа")
        length = _lotki_length(item)
        sheet.append([index, item["startText"], item["endText"], ", ".join(sides), length, round(length * 100), "Выполнен" if item.get("completed") else "Не выполнен"])
    for sheet_item in workbook.worksheets:
        sheet_item.freeze_panes = "A2"
        sheet_item.auto_filter.ref = sheet_item.dimensions
        for column in sheet_item.columns:
            column_letter = column[0].column_letter
            sheet_item.column_dimensions[column_letter].width = min(35, max(12, max(len(str(cell.value or "")) for cell in column) + 2))
    output = BytesIO()
    workbook.save(output)
    return output.getvalue()


def _lotki_export_word(data):
    segments = data.get("segments", [])
    total = sum(_lotki_length(item) for item in segments)
    done = sum(_lotki_length(item) for item in segments if item.get("completed"))
    percent = done / total * 100 if total else 0
    document = Document()
    section = document.sections[0]
    section.top_margin = section.bottom_margin = Cm(1.8)
    document.styles["Normal"].font.name = "Calibri"
    document.styles["Normal"].font.size = Pt(10)
    document.add_heading("Отчёт по лоткам", 0)
    document.add_paragraph(f"Дата формирования: {datetime.now():%d.%m.%Y %H:%M}")
    document.add_paragraph(f"Отрезков: {len(segments)}. Общая длина: {total:.2f} ПК ({total * 100:.0f} м). Выполнено: {percent:.1f}%.")
    table = document.add_table(rows=1, cols=7)
    table.style = "Table Grid"
    headers = ["№", "От ПК", "До ПК", "Сторона", "Длина, ПК", "Длина, м", "Статус"]
    for cell, value in zip(table.rows[0].cells, headers): cell.text = value
    for index, item in enumerate(segments, 1):
        sides = ", ".join((["Слева"] if "left" in item.get("sides", []) else []) + (["Справа"] if "right" in item.get("sides", []) else []))
        length = _lotki_length(item)
        values = [index, item["startText"], item["endText"], sides, f"{length:.2f}", f"{length * 100:.0f}", "Выполнен" if item.get("completed") else "Не выполнен"]
        cells = table.add_row().cells
        for cell, value in zip(cells, values): cell.text = str(value)
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _backup_current_data():
    if not DATA_FILE.is_file():
        return
    BACKUP_DIR.mkdir(exist_ok=True)
    # Автосохранение вызывается часто; используем секунды и миллисекунды, чтобы не терять промежуточные копии.
    stamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S_%f")[:-3]
    target = BACKUP_DIR / f"road-progress-{stamp}.json"
    if target.exists():
        return
    target.write_bytes(DATA_FILE.read_bytes())
    backups = sorted(BACKUP_DIR.glob("*.json"), key=lambda p: p.stat().st_mtime, reverse=True)
    for old in backups[MAX_BACKUPS:]:
        try:
            old.unlink()
        except OSError:
            pass


def _export_excel(data):
    state = _project_state(data)
    workbook = Workbook()
    summary = workbook.active
    summary.title = "Сводка"
    summary.append(["Объект", "ПК от", "ПК до", "Длина слоя, ПК", "Исключено, ПК", "Слой", "Готово, ПК", "План, %", "Дата плана", "Факт, %", "Отклонение, п.п."])
    project_name = next((p.get("name") for p in data.get("projects", []) if p.get("id") == data.get("activeProjectId")), "Основной объект")
    for layer in state.get("layers", []):
        covered = sum(max(0, s["e"] - s["s"]) for s in layer.get("segments", []))
        start = float(state.get("start", 0))
        end = float(state["total"])
        route_length = max(0, end - start)
        excluded = min(route_length, max(0, float(layer.get("excludedPk", 0) or 0)))
        layer_length = max(0, route_length - excluded)
        actual = round(covered / layer_length * 100, 1) if layer_length else 0
        plan = float(layer.get("planPercent", 0))
        summary.append([project_name, start, end, route_length, excluded, layer.get("name"), covered, plan, layer.get("planDate", ""), actual, actual - plan])
    segments = workbook.create_sheet("Участки")
    segments.append(["Слой", "От, ПК", "До, ПК", "Длина, ПК", "Статус", "Ответственный", "Дата", "Комментарий", "Качество"])
    for layer in state.get("layers", []):
        for seg in layer.get("segments", []):
            segments.append([layer.get("name"), seg["s"], seg["e"], seg["e"] - seg["s"], seg.get("status", "Выполнено"), seg.get("responsible", ""), seg.get("date", ""), seg.get("note", ""), seg.get("quality", "")])
    for sheet in workbook.worksheets:
        sheet.freeze_panes = "A2"
        sheet.auto_filter.ref = sheet.dimensions
        for col in sheet.columns:
            sheet.column_dimensions[col[0].column_letter].width = min(45, max(12, max(len(str(c.value or "")) for c in col) + 2))
    output = BytesIO()
    workbook.save(output)
    return output.getvalue()


def _import_excel(raw):
    book = load_workbook(BytesIO(raw), data_only=True)
    sheet = book["Участки"] if "Участки" in book.sheetnames else book.active
    rows = list(sheet.iter_rows(values_only=True))
    if not rows:
        raise ValueError("Excel-файл пуст")
    headers = {str(value or "").strip().lower(): index for index, value in enumerate(rows[0])}
    required = ["слой", "от, пк", "до, пк"]
    if not all(key in headers for key in required):
        raise ValueError("нужны колонки: Слой, От, ПК, До, ПК (лист «Участки»)")
    imported = []
    for row in rows[1:]:
        name = str(row[headers["слой"]] or "").strip()
        if not name:
            continue
        try:
            seg = {"s": float(row[headers["от, пк"]]), "e": float(row[headers["до, пк"]])}
        except (ValueError, TypeError):
            raise ValueError(f"некорректный участок слоя «{name}»")
        for field, header in (("status", "статус"), ("responsible", "ответственный"), ("date", "дата"), ("note", "комментарий"), ("quality", "качество")):
            if header in headers and row[headers[header]] is not None:
                seg[field] = str(row[headers[header]])
        imported.append((name, seg))
    return imported


class Handler(SimpleHTTPRequestHandler):
    def __init__(self, *args, **kwargs):
        super().__init__(*args, directory=str(DIR), **kwargs)

    def do_GET(self):
        path = self.path.split("?", 1)[0]
        if path == "/api/lotki-data":
            self._send_lotki_data()
            return
        if path == "/api/data":
            self._send_data()
            return
        if path == "/api/backups":
            self._send_backups()
            return
        if path == "/api/export/excel":
            self._export_excel()
            return
        if path == "/api/lotki/export/excel":
            self._export_lotki_excel()
            return
        if path == "/favicon.ico":
            self.send_response(204)
            self.end_headers()
            return
        super().do_GET()

    def do_POST(self):
        path = self.path.split("?", 1)[0]
        if path == "/api/lotki-data":
            self._save_lotki_data()
            return
        if path == "/api/data":
            self._save_data()
            return
        if path == "/api/export/word":
            self._export_word()
            return
        if path == "/api/lotki/export/word":
            self._export_lotki_word()
            return
        if path == "/api/import/excel":
            self._import_excel()
            return
        if path == "/api/backups/create":
            self._create_backup()
            return
        if path == "/api/backups/restore":
            self._restore_backup()
            return
        self.send_error(404)

    def _send_data(self):
        if DATA_FILE.is_file():
            body = DATA_FILE.read_bytes()
        else:
            body = b"{}"
        self.send_response(200)
        self.send_header("Content-Type", "application/json; charset=utf-8")
        self.send_header("Cache-Control", "no-store")
        self.send_header("Content-Length", str(len(body)))
        self.end_headers()
        self.wfile.write(body)

    def _send_lotki_data(self):
        body = LOTKI_DATA_FILE.read_bytes() if LOTKI_DATA_FILE.is_file() else b'{"segments":[]}'
        self.send_response(200)
        self.send_header("Content-Type", "application/json; charset=utf-8")
        self.send_header("Cache-Control", "no-store")
        self.send_header("Content-Length", str(len(body)))
        self.end_headers()
        self.wfile.write(body)

    def _save_lotki_data(self):
        length = int(self.headers.get("Content-Length", 0))
        try:
            data = json.loads(self.rfile.read(length).decode("utf-8"))
            _validate_lotki(data)
            _write_json_file(LOTKI_DATA_FILE, data)
        except Exception as exc:
            self.send_error(400, str(exc))
            return
        body = b'{"ok":true}'
        self.send_response(200)
        self.send_header("Content-Type", "application/json")
        self.send_header("Content-Length", str(len(body)))
        self.end_headers()
        self.wfile.write(body)

    def _save_data(self):
        length = int(self.headers.get("Content-Length", 0))
        raw = self.rfile.read(length)
        try:
            data = json.loads(raw.decode("utf-8"))
            _validate_data(data)
            _backup_current_data()
            _write_json_file(DATA_FILE, data)
        except Exception as exc:
            self.send_error(400, str(exc))
            return

        body = b'{"ok":true}'
        self.send_response(200)
        self.send_header("Content-Type", "application/json")
        self.send_header("Content-Length", str(len(body)))
        self.end_headers()
        self.wfile.write(body)

    def _export_word(self):
        length = int(self.headers.get("Content-Length", 0))
        raw = self.rfile.read(length)
        try:
            data = json.loads(raw.decode("utf-8"))
            _validate_data(data)
            body = build_word_document(data)
        except Exception as exc:
            self.send_error(400, str(exc))
            return

        filename = "otchet-progress.docx"
        self.send_response(200)
        self.send_header("Content-Type", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        self.send_header("Content-Disposition", f'attachment; filename="{filename}"')
        self.send_header("Content-Length", str(len(body)))
        self.end_headers()
        self.wfile.write(body)

    def _export_excel(self):
        try:
            data = json.loads(DATA_FILE.read_text(encoding="utf-8")) if DATA_FILE.is_file() else {}
            _validate_data(data)
            body = _export_excel(data)
        except Exception as exc:
            self.send_error(400, str(exc))
            return
        self.send_response(200)
        self.send_header("Content-Type", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
        self.send_header("Content-Disposition", 'attachment; filename="progress.xlsx"')
        self.send_header("Content-Length", str(len(body)))
        self.end_headers()
        self.wfile.write(body)

    def _export_lotki_excel(self):
        try:
            data = json.loads(LOTKI_DATA_FILE.read_text(encoding="utf-8")) if LOTKI_DATA_FILE.is_file() else {"segments": []}
            _validate_lotki(data)
            body = _lotki_export_excel(data)
        except Exception as exc:
            self.send_error(400, str(exc))
            return
        self.send_response(200)
        self.send_header("Content-Type", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
        self.send_header("Content-Disposition", 'attachment; filename="lotki.xlsx"')
        self.send_header("Content-Length", str(len(body)))
        self.end_headers()
        self.wfile.write(body)

    def _export_lotki_word(self):
        length = int(self.headers.get("Content-Length", 0))
        try:
            data = json.loads(self.rfile.read(length).decode("utf-8"))
            _validate_lotki(data)
            body = _lotki_export_word(data)
        except Exception as exc:
            self.send_error(400, str(exc))
            return
        self.send_response(200)
        self.send_header("Content-Type", "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        self.send_header("Content-Disposition", 'attachment; filename="lotki.docx"')
        self.send_header("Content-Length", str(len(body)))
        self.end_headers()
        self.wfile.write(body)

    def _import_excel(self):
        length = int(self.headers.get("Content-Length", 0))
        raw = self.rfile.read(length)
        try:
            imported = _import_excel(raw)
            data = json.loads(DATA_FILE.read_text(encoding="utf-8")) if DATA_FILE.is_file() else {}
            state = _project_state(data)
            layers = {layer["name"]: layer for layer in state.get("layers", [])}
            for name, segment in imported:
                if name not in layers:
                    layers[name] = {"id": "import_" + str(len(layers)), "name": name, "color": "#6BC48C", "segments": []}
                    state.setdefault("layers", []).append(layers[name])
                layers[name]["segments"].append(segment)
            _validate_data(data)
            _backup_current_data()
            _write_json_file(DATA_FILE, data)
            body = json.dumps(data, ensure_ascii=False).encode("utf-8")
        except Exception as exc:
            self.send_error(400, str(exc))
            return
        self.send_response(200)
        self.send_header("Content-Type", "application/json; charset=utf-8")
        self.send_header("Content-Length", str(len(body)))
        self.end_headers()
        self.wfile.write(body)

    def _send_backups(self):
        BACKUP_DIR.mkdir(exist_ok=True)
        backups = [{"name": p.name, "size": p.stat().st_size, "modified": datetime.fromtimestamp(p.stat().st_mtime).isoformat()} for p in sorted(BACKUP_DIR.glob("*.json"), reverse=True)]
        body = json.dumps(backups, ensure_ascii=False).encode("utf-8")
        self.send_response(200)
        self.send_header("Content-Type", "application/json; charset=utf-8")
        self.send_header("Content-Length", str(len(body)))
        self.end_headers()
        self.wfile.write(body)

    def _create_backup(self):
        try:
            _backup_current_data()
            body = b'{"ok":true}'
        except Exception as exc:
            self.send_error(400, str(exc))
            return
        self.send_response(200)
        self.send_header("Content-Type", "application/json")
        self.send_header("Content-Length", str(len(body)))
        self.end_headers()
        self.wfile.write(body)

    def _restore_backup(self):
        length = int(self.headers.get("Content-Length", 0))
        try:
            request = json.loads(self.rfile.read(length).decode("utf-8"))
            name = Path(request["name"]).name
            source = BACKUP_DIR / name
            if not source.is_file() or source.suffix != ".json":
                raise ValueError("резервная копия не найдена")
            data = json.loads(source.read_text(encoding="utf-8"))
            _validate_data(data)
            _backup_current_data()
            _write_json_file(DATA_FILE, data)
            body = json.dumps(data, ensure_ascii=False).encode("utf-8")
        except Exception as exc:
            self.send_error(400, str(exc))
            return
        self.send_response(200)
        self.send_header("Content-Type", "application/json; charset=utf-8")
        self.send_header("Content-Length", str(len(body)))
        self.end_headers()
        self.wfile.write(body)

    def log_message(self, format, *args):
        if args:
            req = str(args[0])
            if req.startswith("GET /api/data") or req.startswith("POST /api/data"):
                return
            if req.startswith("POST /api/export/word"):
                return
            if req.startswith("GET /favicon.ico"):
                return
            if len(args) >= 2 and str(args[1]) == "304":
                return
        super().log_message(format, *args)


def main():
    url = f"http://127.0.0.1:{PORT}/road-progress.html"
    print("Прогресс укладки слоёв")
    print(f"  Страница: {url}")
    print(f"  Данные:   {DATA_FILE.name}")
    print("  Закрой это окно, чтобы остановить сервер.\n")

    try:
        webbrowser.open(url)
    except OSError:
        pass

    server = HTTPServer(("127.0.0.1", PORT), Handler)
    try:
        server.serve_forever()
    except KeyboardInterrupt:
        print("\nСервер остановлен.")
        server.server_close()


if __name__ == "__main__":
    main()
    sys.exit(0)
